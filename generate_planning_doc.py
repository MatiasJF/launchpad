#!/usr/bin/env python3
"""Generate BSVA-branded strategic planning questions document in Spanish."""

import subprocess, sys
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "lxml"])
    from docx import Document

import copy
from lxml import etree
from docx.oxml.ns import qn

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TEMPLATE_PATH = "/Users/matiasjackson/.claude/skills/bsva-docx/BSV Association Word template 2025.docx"

def _set_cnf(parent, val, **kwargs):
    cnf = etree.SubElement(parent, qn("w:cnfStyle"))
    cnf.set(qn("w:val"), val)
    for k, v in kwargs.items():
        cnf.set(qn(f"w:{k}"), v)

def make_paragraph(style_id, text):
    p = etree.SubElement(etree.Element("dummy"), qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), style_id)
    if text:
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = text
        if text and (text[0] == " " or text[-1] == " "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p

def generate_document(output_path):
    doc = Document(TEMPLATE_PATH)
    body = doc.element.body
    children = list(body)

    # Step 1: Modify cover page text
    cover_texts = {
        0: "Protocolo de Financiamiento\nDescentralizado BSV",
        1: "Preguntas Estratégicas de Diseño",
        2: "Versión 1.0 — Planificación",
        3: "Agosto 2026",
    }
    for idx, new_text in cover_texts.items():
        para = children[idx]
        for r in para.findall(f"{{{W}}}r"):
            para.remove(r)
        run = etree.SubElement(para, qn("w:r"))
        if "\n" in new_text:
            parts = new_text.split("\n")
            t = etree.SubElement(run, qn("w:t"))
            t.text = parts[0]
            for part in parts[1:]:
                etree.SubElement(run, qn("w:br"))
                t2 = etree.SubElement(run, qn("w:t"))
                t2.text = part
        else:
            t = etree.SubElement(run, qn("w:t"))
            t.text = new_text

    # Step 2: Remove sample content
    to_remove = [children[i] for i in range(4, 90)] + [children[91]]
    for child in to_remove:
        body.remove(child)

    # Step 3: Merge sections
    children = list(body)
    sectpr_para = children[4]
    embedded_sectpr = sectpr_para.find(f".//{{{W}}}sectPr")
    final_sectpr = children[5]
    for child in list(final_sectpr):
        final_sectpr.remove(child)
    for attr, val in embedded_sectpr.attrib.items():
        final_sectpr.set(attr, val)
    for child in list(embedded_sectpr):
        final_sectpr.append(copy.deepcopy(child))
    body.remove(sectpr_para)

    # Step 4: Update header/footer
    for section in doc.sections:
        if not section.header.is_linked_to_previous:
            for para in section.header.paragraphs:
                for run in para.runs:
                    if "BSV Association Document Template" in (run.text or ""):
                        run.text = "Protocolo de Financiamiento Descentralizado BSV"
        if not section.footer.is_linked_to_previous:
            for para in section.footer.paragraphs:
                for run in para.runs:
                    if "Version 1.1" in (run.text or ""):
                        run.text = "Versión 1.0 — Planificación"

    # Step 5: Insert content
    final_sectpr = list(body)[-1]
    content = [
        make_paragraph("BSVAIntroduction", "Este documento presenta las preguntas estratégicas críticas que deben responderse antes de diseñar un protocolo de financiamiento colectivo (crowdfunding) completamente descentralizado sobre la blockchain de BSV. Las respuestas determinarán la arquitectura técnica, el modelo de confianza, y la viabilidad del sistema."),
        make_paragraph("Normal", ""),

        make_paragraph("Heading1", "1. Modelo de Confianza — ¿Qué DEBE estar en cadena vs qué PUEDE estar fuera de cadena?"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "¿Cuáles de las siguientes operaciones DEBEN ser aplicadas por la blockchain (covenant/script), y cuáles son aceptables de rastrear fuera de cadena con verificación SPV?"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet", "Creación de proyectos (metadata: nombre, descripción, meta de financiamiento)"),
        make_paragraph("ListBullet", "Reglas de financiamiento (quién puede contribuir, montos mín/máx, fecha límite)"),
        make_paragraph("ListBullet", "Custodia de fondos (dónde se mantienen los fondos recaudados durante la campaña)"),
        make_paragraph("ListBullet", "Lógica de reembolso (reembolso automático si no se alcanza la meta, o retiro de emergencia)"),
        make_paragraph("ListBullet", "Distribución de tokens (quién recibe qué tokens, a qué precio)"),
        make_paragraph("ListBullet", "Liberación por hitos (liberación gradual de fondos vinculada a entregables)"),
        make_paragraph("Normal", ""),
        make_paragraph("BSVAQuote", "Pregunta concreta: Si el operador/sitio web desaparece a mitad de campaña, ¿qué operaciones DEBE poder ejecutar un contribuyente con SOLO la blockchain y su billetera?"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet", "¿Demostrar que contribuyó?"),
        make_paragraph("ListBullet", "¿Obtener un reembolso si no se alcanzó la meta?"),
        make_paragraph("ListBullet", "¿Reclamar tokens si se alcanzó la meta?"),
        make_paragraph("ListBullet", "¿Verificar que el propietario del proyecto no puede hacer un 'rug pull'?"),
        make_paragraph("Normal", ""),

        make_paragraph("Heading1", "2. Curva de Vinculación (Bonding Curve) — ¿Requerida u Opcional?"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "La curva de vinculación (el precio aumenta con la oferta) es UN mecanismo de financiamiento, pero no el único."),
        make_paragraph("Normal", ""),
        make_paragraph("BSVAQuote", "Pregunta: ¿Es la curva de vinculación (precio dinámico) esencial para su visión, o es UNA implementación entre varios modelos de financiamiento aceptables?"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "Modelos alternativos que evitan el cuello de botella del UTXO único:"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet", "Preventa a precio fijo (todos pagan lo mismo, primero en llegar hasta el límite) — ya funciona en instant_swap"),
        make_paragraph("ListBullet", "Subasta holandesa (el precio baja con el tiempo hasta que se agota)"),
        make_paragraph("ListBullet", "Subasta por lotes (recopilar todas las ofertas, liquidar al precio de mercado una vez)"),
        make_paragraph("ListBullet", "Vesting vinculado a hitos (los fondos se desbloquean a medida que se demuestran los hitos)"),
        make_paragraph("Normal", ""),

        make_paragraph("Heading1", "3. Requisitos de Escala — ¿Cuántos proyectos, cuántos contribuyentes por proyecto?"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "Esto determina si necesitamos optimizar para:"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet", "Muchos proyectos, pocos contribuyentes cada uno (crowdfunding de nicho) → aislamiento de proyectos crítico"),
        make_paragraph("ListBullet", "Pocos proyectos, muchos contribuyentes cada uno (campañas virales) → alto rendimiento por proyecto crítico"),
        make_paragraph("Normal", ""),
        make_paragraph("BSVAQuote", "Pregunta: ¿Cuál es un escenario de éxito realista en el año 1?"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet", "10 proyectos × 100 contribuyentes cada uno = 1,000 contribuciones totales?"),
        make_paragraph("ListBullet", "100 proyectos × 1,000 contribuyentes cada uno = 100,000 contribuciones totales?"),
        make_paragraph("ListBullet", "¿Algo diferente?"),
        make_paragraph("Normal", ""),

        make_paragraph("Heading1", "4. Regulatorio / KYC — ¿Necesitamos participación con permisos?"),
        make_paragraph("Normal", ""),
        make_paragraph("BSVAQuote", "Pregunta: ¿Puede cualquiera contribuir a cualquier proyecto (sin permisos), o necesitamos:"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet", "Aprobación del propietario del proyecto por contribuyente (lista blanca)?"),
        make_paragraph("ListBullet", "Compuertas KYC/AML antes de contribuciones grandes?"),
        make_paragraph("ListBullet", "Restricciones geográficas?"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "Esto afecta si el covenant puede ser completamente sin permisos o necesita una puerta de operador."),
        make_paragraph("Normal", ""),

        make_paragraph("Heading1", "5. Independencia del Frontend — ¿Qué significa \"protocolo\" aquí?"),
        make_paragraph("Normal", ""),
        make_paragraph("BSVAQuote", "Pregunta: Cuando dice 'cualquiera puede construir un frontend sobre él', ¿qué exactamente debe ser interoperable?"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet", "Nivel 1 (Solo lectura): Cualquier frontend puede MOSTRAR proyectos/contribuciones leyendo la cadena"),
        make_paragraph("ListBullet", "Nivel 2 (Contribuir): Cualquier frontend puede permitir a los usuarios CONTRIBUIR a proyectos existentes"),
        make_paragraph("ListBullet", "Nivel 3 (Crear): Cualquier frontend puede permitir a los usuarios CREAR nuevos proyectos"),
        make_paragraph("ListBullet", "Nivel 4 (Liquidar): Cualquier frontend puede ACTIVAR reembolsos/distribuciones (no se necesita operador)"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "¿Cuál nivel es el objetivo?"),
        make_paragraph("Normal", ""),

        make_paragraph("Heading1", "Próximos Pasos"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "Una vez respondidas estas preguntas, se realizará una investigación exhaustiva y orquestada en paralelo sobre:"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet2", "Capacidades de covenants de BSV (qué puede y no puede hacer Script)"),
        make_paragraph("ListBullet2", "Soluciones alternativas para UTXO (fragmentación, agregación fuera de cadena, modelos híbridos)"),
        make_paragraph("ListBullet2", "Protocolos BSV existentes (redes de superposición, estándares de tokens, patrones de covenant)"),
        make_paragraph("ListBullet2", "Arquitecturas alternativas (liquidación por lotes, pools de liquidez, árboles de Merkle)"),
        make_paragraph("ListBullet2", "Proyectos comparables (qué hace el crowdfunding de Ethereum/Cosmos, adaptado a UTXO)"),
        make_paragraph("Normal", ""),
        make_paragraph("Normal", "El resultado será un documento de estrategia con marca BSVA que incluirá:"),
        make_paragraph("Normal", ""),
        make_paragraph("ListBullet2", "Diagramas de arquitectura Mermaid (límites de confianza, flujos de transacciones, máquinas de estado)"),
        make_paragraph("ListBullet2", "Análisis de viabilidad (qué es posible hoy vs qué necesita I+D)"),
        make_paragraph("ListBullet2", "Camino recomendado hacia adelante (hoja de ruta por fases)"),
        make_paragraph("ListBullet2", "Evaluación de riesgos (qué se rompe si el operador desaparece)"),
        make_paragraph("Normal", ""),
    ]

    for elem in content:
        body.insert(list(body).index(final_sectpr), elem)

    doc.save(output_path)
    print(f"✓ Documento generado: {output_path}")

if __name__ == "__main__":
    output = "/Users/matiasjackson/Documents/Proyects/exchanges_listings/launchpad/Preguntas_Estrategicas_Protocolo_BSV.docx"
    generate_document(output)
