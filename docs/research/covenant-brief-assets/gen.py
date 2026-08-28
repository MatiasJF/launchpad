#!/usr/bin/env python3
"""BSVA-branded covenant audit brief."""
import subprocess, sys, copy
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "lxml", "-q"])
    from docx import Document
from docx.shared import Inches
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TEMPLATE_PATH = "/Users/matiasjackson/.claude/skills/bsva-docx/BSV Association Word template 2025.docx"
HERE = os.path.dirname(os.path.abspath(__file__))
IMG = HERE                       # the .png files sit beside this script
OUT = os.path.join(HERE, "..", "BSVA-Covenant-Review-Brief.docx")

def make_paragraph(style_id, text):
    p = etree.SubElement(etree.Element("dummy"), qn("w:p"))
    pPr = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
    pStyle.set(qn("w:val"), style_id)
    if text:
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = text
        if text[0] == " " or text[-1] == " ":
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p

def _set_cnf(parent, val, **kwargs):
    cnf = etree.SubElement(parent, qn("w:cnfStyle")); cnf.set(qn("w:val"), val)
    for k, v in kwargs.items(): cnf.set(qn(f"w:{k}"), v)

def _add_cell(tr, text, w, style_id, first_col=False, shaded=False, vAlign=None):
    tc = etree.SubElement(tr, qn("w:tc")); tcPr = etree.SubElement(tc, qn("w:tcPr"))
    if first_col: _set_cnf(tcPr, "001000000000", firstColumn="1")
    tw_el = etree.SubElement(tcPr, qn("w:tcW")); tw_el.set(qn("w:w"), str(w)); tw_el.set(qn("w:type"), "pct")
    if shaded:
        shd = etree.SubElement(tcPr, qn("w:shd"))
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "EFF0F7"); shd.set(qn("w:themeFill"), "background2")
    if vAlign:
        va = etree.SubElement(tcPr, qn("w:vAlign")); va.set(qn("w:val"), vAlign)
    p = etree.SubElement(tc, qn("w:p")); pPr = etree.SubElement(p, qn("w:pPr"))
    pStyle = etree.SubElement(pPr, qn("w:pStyle")); pStyle.set(qn("w:val"), style_id)
    r = etree.SubElement(p, qn("w:r")); t = etree.SubElement(r, qn("w:t")); t.text = str(text)
    if text and (str(text)[0] == " " or str(text)[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

def make_table_xml(headers, data_rows, widths=None):
    ncols = len(headers)
    if widths: col_widths = widths
    else:
        base_w = 5000 // ncols
        col_widths = [base_w] * ncols; col_widths[-1] = 5000 - base_w * (ncols - 1)
    tbl = etree.SubElement(etree.Element("dummy"), qn("w:tbl"))
    tblPr = etree.SubElement(tbl, qn("w:tblPr"))
    etree.SubElement(tblPr, qn("w:tblStyle")).set(qn("w:val"), "nChain")
    tw = etree.SubElement(tblPr, qn("w:tblW")); tw.set(qn("w:w"), "5000"); tw.set(qn("w:type"), "pct")
    tl = etree.SubElement(tblPr, qn("w:tblLook"))
    for k, v in {"val":"04A0","firstRow":"1","lastRow":"0","firstColumn":"1","lastColumn":"0","noHBand":"0","noVBand":"1"}.items():
        tl.set(qn(f"w:{k}"), v)
    tblGrid = etree.SubElement(tbl, qn("w:tblGrid"))
    for w in col_widths:
        gc = etree.SubElement(tblGrid, qn("w:gridCol")); gc.set(qn("w:w"), str(int(w * 9864 / 5000)))
    tr = etree.SubElement(tbl, qn("w:tr")); trPr = etree.SubElement(tr, qn("w:trPr"))
    _set_cnf(trPr, "100000000000", firstRow="1")
    etree.SubElement(trPr, qn("w:trHeight")).set(qn("w:val"), "432")
    for ci, (text, w) in enumerate(zip(headers, col_widths)):
        _add_cell(tr, text, w, "BSVATableHeader", first_col=(ci==0), vAlign="center")
    for ri, row_data in enumerate(data_rows):
        shaded = (ri % 2 == 0)
        tr = etree.SubElement(tbl, qn("w:tr")); trPr = etree.SubElement(tr, qn("w:trPr"))
        if shaded: _set_cnf(trPr, "000000100000", oddHBand="1")
        etree.SubElement(trPr, qn("w:trHeight")).set(qn("w:val"), "455")
        for ci, (text, w) in enumerate(zip(row_data, col_widths)):
            _add_cell(tr, text, w, "BSVATableText", first_col=(ci==0), shaded=shaded)
    return tbl

def P(t): return make_paragraph("Normal", t)
def H1(t): return make_paragraph("Heading1", t)
def H2(t): return make_paragraph("Heading2", t)
def B(t): return make_paragraph("ListBullet", t)
def Q(t): return make_paragraph("BSVAQuote", t)
def INTRO(t): return make_paragraph("BSVAIntroduction", t)
def FIG(name, w): return make_paragraph("Normal", f"@@IMG:{name}:{w}")
def CAP(t): return make_paragraph("Normal", t)

def build():
    doc = Document(TEMPLATE_PATH); body = doc.element.body; children = list(body)
    cover = {0: "Trustless Bonding-Curve Covenant",
             1: "Technical review brief",
             2: "Version 1.0  ·  August 2026  ·  Internal",
             3: ""}
    for idx, txt in cover.items():
        para = children[idx]
        for r in para.findall(f"{{{W}}}r"): para.remove(r)
        run = etree.SubElement(para, qn("w:r")); t = etree.SubElement(run, qn("w:t")); t.text = txt
    for child in [children[i] for i in range(4, 90)] + [children[91]]: body.remove(child)
    children = list(body)
    sectpr_para = children[4]; embedded = sectpr_para.find(f".//{{{W}}}sectPr"); final = children[5]
    for c in list(final): final.remove(c)
    for a, v in embedded.attrib.items(): final.set(a, v)
    for c in list(embedded): final.append(copy.deepcopy(c))
    body.remove(sectpr_para)
    for section in doc.sections:
        if not section.header.is_linked_to_previous:
            for para in section.header.paragraphs:
                for run in para.runs:
                    if "BSV Association Document Template" in (run.text or ""):
                        run.text = "Trustless Bonding-Curve Covenant — Review Brief"
        if not section.footer.is_linked_to_previous:
            for para in section.footer.paragraphs:
                for run in para.runs:
                    if "Version 1.1" in (run.text or ""): run.text = "Version 1.0"

    C = []
    C += [H1("What we are asking you to do")]
    C += [INTRO("Review one Bitcoin Script covenant that custodies user money, and tell us whether anyone can take satoshis they are not owed, or lock satoshis so that nobody can ever retrieve them.")]
    C += [P("That is the whole question. Everything else in this document exists to get you to it faster."), P("")]
    C += [P("In scope — the covenant and the code that constructs its spends:")]
    C += [B("packages/curve/src/contracts/merkleLedgerPool.ts — the deployed covenant (sCrypt source)"),
          B("packages/curve/src/merkleLedger.ts — the off-chain tree that must agree with the in-script fold"),
          B("packages/curve/service/merkleLedgerState.ts — successor-script and unlock construction"),
          B("packages/curve/src/covenant.ts — the pre-broadcast interpreter guard")]
    C += [P(""), P("Out of scope, and we would rather you did not spend time on them: the web application, code style, architecture preferences, and the database layer. None of them can move funds — the covenant is the sole authority, and a bug in the app produces a rejected transaction, not a loss.")]

    C += [H1("What holds the money")]
    C += [P("A single covenant output holds the entire reserve of a sale. Buyers pay into it; sellers are refunded out of it; at the end its whole balance is released to one address that was fixed when the pool was deployed."), P("")]
    C += [Q("There is no operator key anywhere on this path. We cannot stop, redirect, reprice or seize a trade, and neither can the project. If a satoshi can be stolen, it is because the script permits it.")]
    C += [P(""), FIG("fig1.png", 5.0), CAP("Figure 1 — Every satoshi sits in the covenant. The three arrows out are the only ways it can leave.")]

    C += [H1("How it can be spent")]
    C += [P("Three public methods, each with a different sighash discipline. The flags matter as much as the logic: they decide which outputs the covenant commits to, and therefore what an attacker is free to change.")]
    C += [FIG("fig2.png", 6.3), CAP("Figure 2 — The three spend paths and what each one pins.")]
    C += [P(""), P("Buy is keyless by design: nobody signs it. The covenant computes the curve price itself and requires the reserve to grow by at least that amount, so there is no authorising key to compromise. Sell is authorised by the holder's own signature over the spend — that signature is their claim to the balance. Graduation carries no signature at all; anyone may trigger it once the curve is fully sold, and the destination was fixed at deploy so a hostile caller cannot steer the money.")]

    C += [H1("The mechanism most worth attacking")]
    C += [P("Holder balances live inside the covenant as a Merkle root over a fixed-depth array of slots. This is the highest-value target in the system: if a spender can prove a balance they do not hold, they can sell it and drain the reserve.")]
    C += [FIG("fig3.png", 5.0), CAP("Figure 3 — Every spend proves the current value of the slot it touches.")]
    C += [P(""), P("Depth is 16, giving 65,536 slots and a 512-byte inclusion proof that does not grow with the holder count. A leaf is sha256(ownerPkh(20) || balance(8, little-endian)); both fields are fixed-width, so no two distinct pairs can serialise identically. Slots are append-only: a new holder must prove the slot at exactly holderCount is empty, and an existing one must prove the slot's current owner and balance. Nothing can be reset.")]

    C += [H1("Invariants to verify")]
    C += [make_table_xml(["#", "Invariant", "Why it matters"], [
        ["1", "sold equals the sum of all slot balances, always", "The reserve-safety invariant. Every buy adds to both; every sell subtracts from both."],
        ["2", "The curve division never truncates", "d(2s+d+1) is always even, so there is no rounding in anyone's favour. Confirm, do not assume."],
        ["3", "The in-script fold equals the off-chain fold", "Same sibling order, depth and hash. A disagreement means forged balances validate. Sharpest edge here."],
        ["4", "A new holder can only append onto a proven-empty slot", "Otherwise a live balance is overwritten and invariant 1 breaks."],
        ["5", "An update proves the slot's current owner and balance", "The leaf commits to the owner, so a slot cannot be reset or written over by another party."],
        ["6", "An empty slot and an occupied slot are distinguishable", "Empty is 32 zero bytes; a real leaf is a sha256 image."],
        ["7", "hashOutputs pinning matches the flag on each method", "Buy and graduate pin output 0 only; sell pins exactly two outputs."],
        ["8", "Successor scripts are byte-exact", "A one-byte drift fails hashOutputs. This trap has been hit twice in this codebase."],
    ], widths=[400, 1900, 2700])]

    C += [H1("Drain vectors, ranked")]
    C += [P("Attack these in order. The first is the reason the covenant exists.")]
    C += [make_table_xml(["Rank", "Vector", "What to try"], [
        ["1", "Merkle proof forgery", "Wrong sibling order, an off-by-one in depth, a path-bit mismatch — anything that proves a balance not held."],
        ["2", "Append onto an occupied slot", "Defeat either half of the guard independently: the index check, or the empty-leaf proof."],
        ["3", "Cross-owner slot write", "Credit or debit a slot whose owner differs from the one supplied."],
        ["4", "Stale-proof replay", "Reuse a proof from an earlier root."],
        ["5", "Reserve underflow on a sell", "Make the refund exceed the reserve."],
        ["6", "OP_PUSH_TX preimage malleability", "Verify the preimage is validated and that ctx.utxo.value cannot be spoofed."],
        ["7", "holderCount manipulation", "Increment without an append, or append without an increment."],
    ], widths=[600, 1700, 2700])]
    C += [P(""), P("Not applicable here, unlike our other curve: there is no operator key, so no key-compromise drain, and no off-chain gate to defeat.")]
    return doc, body, C

def finish(doc, body, C, out):
    C += [H1("What is already proven on mainnet")]
    C += [P("All of the following ran on BSV mainnet with real satoshis, not on a testnet or a simulator. Transaction identifiers are given so you can inspect the bytes yourself rather than take our word for it.")]
    C += [make_table_xml(["What", "Result", "Reference"], [
        ["Full lifecycle", "deploy, buy, sell, sell-out, graduate, mint, deliver", "pool 38d331f7"],
        ["Script size is constant in holder count", "11,864 bytes at every step", "pool 4c6faf97"],
        ["Byte-level attacks on the unlocking script", "34 of 34 repelled", "verify-merkle-adversarial.ts"],
        ["Ledger rebuilt from chain alone, no database", "16 of 16, holder recovered whose key no longer exists", "verify-merkle-resolve.ts"],
        ["Permissionless graduation by a stranger", "reserve reached the committed payout; caller net negative", "graduation 82e5dd53"],
        ["Contention: two spenders, same tip", "loser rejected by the node, then re-signed and landed", "pool 31820de7"],
        ["Economic invariants under random buy/sell", "54 unit tests including a 40-seed fuzz", "merkle-solvency.test.mjs"],
    ], widths=[1700, 1900, 1400])]

    C += [H1("What we have not tested")]
    C += [P("Stated plainly, because a review is worth less if it starts from an inflated picture of our coverage.")]
    C += [B("The Script has been attacked but not fuzzed. Thirty-four hand-designed byte-level attacks is not randomised or mutational fuzzing. Untried: random sibling and path corruption at scale, malformed pushdata framing, out-of-range path values, and non-minimal scriptNum encodings of the numeric arguments."),
          B("No formal argument that the off-chain and in-script folds are equivalent. They agree on every case we test. That is not a proof."),
          B("The depth boundary is unexercised live. Slot 65,535 proves correctly off-chain, but genuine slot exhaustion would take 65,536 real appends and has never happened."),
          B("The 8-byte balance ceiling is untested on-chain. Nothing in the covenant bounds k or supply, so a deployer could in principle choose terms a balance cannot represent. The curve cost explodes long before that binds, but it is not enforced."),
          B("Graduation delivery is not enforced. Converting final balances into wallet-held tokens after graduation is a promise, not a covenant guarantee — see the note below.")]

    C += [H1("One known and accepted weakness")]
    C += [P("Everything up to graduation is enforced by the script. The step after it is not.")]
    C += [P("When a curve sells out, the reserve goes to the project and holders are left with balances in a final ledger. Converting those into wallet-held tokens is something the project must choose to do. We cannot make the covenant do it, because a token's ownership gate is a pay-to-public-key-hash check requiring a signature, and a covenant spends by inspecting a transaction preimage rather than by holding keys. A covenant can therefore never custody these tokens, which rules out the obvious design of releasing them against an inclusion proof.")]
    C += [P("What survives is accountability rather than enforcement: the amount owed is recomputable from the deploy transaction forever, by anyone, without the project's cooperation. We surface it in the interface before a buyer commits. We would value your opinion on whether the bond mechanism we have specified to close this is worth building, or whether the disclosure is the right stopping point.")]

    C += [H1("What we would like back")]
    C += [B("A finding-by-finding response to the eight invariants and seven drain vectors above."),
          B("Explicit coverage of the five untested areas — especially fuzzing, which we cannot credibly do on our own work."),
          B("A go or no-go, with a maximum reserve size you would be comfortable with."),
          B("A re-audit trigger list: which changes should send this back to you.")]
    C += [P(""), P("There is no deadline pressure and no commercial dependency on a positive answer. If the covenant is unsound we would rather learn it now, while the only money at risk is roughly ninety thousand satoshis of our own test funds.")]

    C += [H1("Where to start")]
    C += [P("The full audit package, including the material summarised here, is in the repository at docs/AUDIT-PREP-MERKLE-LEDGER.md. It carries the same invariants and vectors with file-level references. The covenant itself is under two hundred lines of sCrypt and is the right place to begin.")]
    C += [P(""), P("A second, separate covenant exists for our operator-settled curve (docs/COVENANT-AUDIT-PREP.md). It has a different trust model — it does have an operator key — so findings do not transfer between the two. Review whichever is being deployed; if that is both, they are two engagements.")]

    final_sectpr = list(body)[-1]
    for elem in C: body.insert(list(body).index(final_sectpr), elem)
    doc.save(out)

    # embed images where sentinels were placed
    doc2 = Document(out)
    for para in doc2.paragraphs:
        if para.text.startswith("@@IMG:"):
            _, name, width = para.text.split(":")
            for r in list(para.runs): r._element.getparent().remove(r._element)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(f"{IMG}/{name}", width=Inches(float(width)))
    doc2.save(out)
    print(f"Saved: {out}")

doc, body, C = build()
finish(doc, body, C, os.path.normpath(OUT))
